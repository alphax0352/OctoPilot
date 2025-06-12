from flask import Flask, request, jsonify
import os
import json
import re
from docx import Document
from docx2pdf import convert
from datetime import datetime
import comtypes.client
import comtypes

app = Flask(__name__)

def sanitize_filename(filename: str, replacement: str = "_") -> str:
    """
    Converts an invalid file name string to a valid one by replacing
    illegal characters with a safe replacement character.
    
    :param filename: The original filename string.
    :param replacement: The character to replace invalid characters with.
    :return: A valid filename string.
    """
    # Define invalid characters (common across Windows, macOS, and Linux)
    invalid_chars = r'[<>:"/\\|?*\x00-\x1F]'
    # Replace invalid characters
    sanitized = re.sub(invalid_chars, replacement, filename)
    # Trim leading/trailing whitespace and dots
    sanitized = sanitized.strip(" .")
    return sanitized

def process_paragraphs(paragraphs, replacement_dict):
    """Process paragraphs and replace placeholders, handling cases where placeholders span multiple runs."""
    for paragraph in paragraphs:
        # Get full paragraph text by joining all runs
        full_text = "".join(run.text for run in paragraph.runs)
        for placeholder, value in replacement_dict.items():
            if placeholder in full_text:
                # Find start and end positions of the placeholder
                placeholder_start = full_text.find(placeholder)
                placeholder_end = placeholder_start + len(placeholder)
                current_pos = 0
                for run in paragraph.runs:
                    run_start = current_pos
                    run_end = current_pos + len(run.text)
                    if run_start <= placeholder_start < run_end and run_end >= placeholder_end:
                        # Placeholder is fully within this run
                        run.text = run.text.replace(placeholder, value)
                    elif run_start <= placeholder_start < run_end:
                        # Placeholder starts in this run
                        run.text = run.text[:placeholder_start - run_start] + value
                    elif run_start < placeholder_end <= run_end:
                        # Placeholder ends in this run
                        run.text = run.text[placeholder_end - run_start:]
                    elif placeholder_start <= run_start and run_end <= placeholder_end:
                        # Run is entirely within the placeholder
                        run.text = ""
                    current_pos = run_end

def get_all_text(doc):
    """Extract all text from the document for validation."""
    text = []
    for paragraph in doc.paragraphs:
        text.append(paragraph.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    text.append(paragraph.text)
    return "\n".join(text)

def validate_replacements(doc, replacement_dict):
    """Check for any remaining placeholders and return warnings."""
    all_text = get_all_text(doc)
    warnings = []
    for placeholder in replacement_dict.keys():
        if placeholder in all_text:
            warnings.append(f"Warning: Placeholder '{placeholder}' was not fully replaced in output.")
    return warnings

@app.route('/api/build_resume', methods=['POST'])
def build_resume():
    try:
        # Get replacement_dict from request JSON
        data = request.get_json()
        # print(data)
        if not data or 'resume_content' not in data or 'job_title' not in data or 'company_name' not in data or 'profile_data' not in data or 'employment_history' not in data or 'education_info' not in data:
            return jsonify({'error': 'details are required in the request body'}), 400
        
        resume_content = json.loads(data['resume_content'])
        job_title = data['job_title']
        company_name = data['company_name']
        profile_data = json.loads(data['profile_data'])
        employment_history = json.loads(data['employment_history'])
        education_info = json.loads(data['education_info'])
        resume_template_path = data['resume_template_path']

        replacement_dict = {
            "[Name]": profile_data['name'],
            "[Headline]": resume_content['headline'],
            "[Contact Information]": " • ".join(filter(None, [profile_data['location'], profile_data['email'], profile_data['phone'], profile_data['linkedin']])),
            "[Summary]": resume_content['summary'],
            "[Skills]": resume_content['skills'],
            "[Session 1]": employment_history[0]['from'] + " - " + employment_history[0]['to'],
            "[Location 1]": employment_history[0]['location'],
            "[Title 1]": employment_history[0]['title'],
            "[Company 1]": employment_history[0]['company'],
            "[Accomplishment Bullet Point 1]": resume_content['bullets_first_company'][0],
            "[Accomplishment Bullet Point 2]": resume_content['bullets_first_company'][1],
            "[Accomplishment Bullet Point 3]": resume_content['bullets_first_company'][2],
            "[Accomplishment Bullet Point 4]": resume_content['bullets_first_company'][3],
            "[Accomplishment Bullet Point 5]": resume_content['bullets_first_company'][4],
            "[Accomplishment Bullet Point 6]": resume_content['bullets_first_company'][5],
            "[Accomplishment Bullet Point 7]": resume_content['bullets_first_company'][6],
            "[Session 2]": employment_history[1]['from'] + " - " + employment_history[1]['to'],
            "[Location 2]": employment_history[1]['location'],
            "[Title 2]": employment_history[1]['title'],
            "[Company 2]": employment_history[1]['company'],
            "[Accomplishment Bullet Point 8]": resume_content['bullets_second_company'][0],
            "[Accomplishment Bullet Point 9]": resume_content['bullets_second_company'][1],
            "[Accomplishment Bullet Point 10]": resume_content['bullets_second_company'][2],
            "[Accomplishment Bullet Point 11]": resume_content['bullets_second_company'][3],
            "[Accomplishment Bullet Point 12]": resume_content['bullets_second_company'][4],
            "[Session 3]": employment_history[2]['from'] + " - " + employment_history[2]['to'],
            "[Location 3]": employment_history[2]['location'],
            "[Title 3]": employment_history[2]['title'],
            "[Company 3]": employment_history[2]['company'],
            "[Accomplishment Bullet Point 13]": resume_content['bullets_third_company'][0],
            "[Accomplishment Bullet Point 14]": resume_content['bullets_third_company'][1],
            "[Accomplishment Bullet Point 15]": resume_content['bullets_third_company'][2],
            "[Session 4]": education_info['from'] + " - " + education_info['to'],
            "[Location 4]": education_info['location'],
            "[Degree]": education_info['degree'],
            "[University]": education_info['school'],
        }
        
        input_docx = resume_template_path

        # Input and output file paths
        script_dir = os.path.dirname(os.path.abspath(__file__))
        datestamp = datetime.now().strftime("%m-%d")
        timestamp = datetime.now().strftime("%H-%M-%S")
        output_filename = f"{sanitize_filename(profile_data['name'])}_{sanitize_filename(job_title)}.docx"
        output_docx = os.path.join(script_dir, "outputs", sanitize_filename(profile_data['email']), datestamp, f"{timestamp}_{sanitize_filename(company_name)}", output_filename)
        output_pdf = os.path.join(script_dir, "outputs", sanitize_filename(profile_data['email']), datestamp, f"{timestamp}_{sanitize_filename(company_name)}", output_filename.replace(".docx", ".pdf"))
        docx_path = os.path.join("resume", "outputs", sanitize_filename(profile_data['email']), datestamp, f"{timestamp}_{sanitize_filename(company_name)}", output_filename)
        pdf_path = os.path.join("resume", "outputs", sanitize_filename(profile_data['email']), datestamp, f"{timestamp}_{sanitize_filename(company_name)}", output_filename.replace(".docx", ".pdf"))
        
        # Create outputs directory if it doesn't exist
        os.makedirs(os.path.dirname(output_docx), exist_ok=True)
        
        # Check if input file exists
        if not os.path.exists(input_docx):
            return jsonify({'error': f"Input file '{input_docx}' does not exist"}), 400

        # Load the input document
        doc = Document(input_docx)

        # Process paragraphs in the main body
        process_paragraphs(doc.paragraphs, replacement_dict)

        # Process paragraphs in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    process_paragraphs(cell.paragraphs, replacement_dict)

        # Save the modified .docx document
        doc.save(output_docx)

        # Initialize COM for this thread
        comtypes.CoInitialize()
        
        try:
            # Try to convert .docx to .pdf with better error handling
            convert(output_docx, output_pdf)
        except Exception as pdf_error:
            print(f"PDF conversion error: {str(pdf_error)}")
            # If PDF conversion fails, still return the DOCX
            response = {
                'docx_path': docx_path,
                'pdf_path': None,
                'warnings': warnings + [f"PDF conversion failed: {str(pdf_error)}"]
            }
            return jsonify(response), 200
        finally:
            # Uninitialize COM to clean up
            comtypes.CoUninitialize()

        # Validate replacements
        output_doc = Document(output_docx)
        warnings = validate_replacements(output_doc, replacement_dict)

        # Return response with output path and any warnings
        response = {
            'docx_path': docx_path,
            'pdf_path': pdf_path,
            'warnings': warnings
        }
        return jsonify(response), 200

    except Exception as e:
        import traceback
        error_details = traceback.format_exc()
        print(f"💥Detailed error: {error_details}")  # This will print to the server console
        return jsonify({
            'error': f"Error processing document: {str(e)}",
            'details': error_details
        }), 500

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=3830, debug=True)